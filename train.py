# # Import necessary libraries
# import numpy as np
# import pandas as pd
# import re
# import nltk
# import pdfplumber
# import os
# from docx import Document  # <--- NEW
# from tensorflow.keras.models import Sequential
# from tensorflow.keras.layers import Embedding, Conv1D, MaxPooling1D, LSTM, Dense, Bidirectional
# from tensorflow.keras.preprocessing.text import Tokenizer
# from tensorflow.keras.preprocessing.sequence import pad_sequences
# from sklearn.model_selection import train_test_split
# from collections import Counter
# from nltk.corpus import stopwords
# from nltk.stem import PorterStemmer

# nltk.download('stopwords')

# # Step 1: Extract Text from PDFs
# def extract_text_from_pdf(pdf_path):
#     text = ""
#     with pdfplumber.open(pdf_path) as pdf:
#         for page in pdf.pages:
#             page_text = page.extract_text()
#             if page_text:
#                 text += page_text + "\n"
#     return text

# # Step 1 (Alternate): Extract Text from DOCX
# def extract_text_from_docx(docx_path):
#     text = ""
#     doc = Document(docx_path)
#     for para in doc.paragraphs:
#         text += para.text + "\n"
#     return text

# # Specify the folder containing your resumes
# resume_folder = r'H:\AI-Resume-Parser-main\Resumes'
#   # <-- Remove extra quotes here
# resume_texts = []
# for filename in os.listdir(resume_folder):
#     file_path = os.path.join(resume_folder, filename)
#     if filename.endswith('.pdf'):
#         text = extract_text_from_pdf(file_path)
#     elif filename.endswith('.docx'):
#         text = extract_text_from_docx(file_path)
#     else:
#         continue
#     resume_texts.append({'filename': filename, 'text': text})

# # Convert to DataFrame
# df = pd.DataFrame(resume_texts)

# # Step 2: Data Preprocessing
# def clean_text(text):
#     text = re.sub(r'[^a-zA-Z\s]', '', text)
#     text = text.lower().strip()
#     return text

# df['cleaned_text'] = df['text'].apply(clean_text)

# # Tokenization and stopword removal
# stop_words = set(stopwords.words('english'))
# df['tokens'] = df['cleaned_text'].apply(lambda x: [word for word in x.split() if word not in stop_words])

# # Stemming
# ps = PorterStemmer()
# df['stemmed'] = df['tokens'].apply(lambda x: [ps.stem(word) for word in x])

# # Label mapping
# keyword_to_label = {
#     'manager': 'Manager',
#     'sales manager': 'Sales Manager',
#     'project manager': 'Project Manager',
#     'analyst': 'Analyst',
#     'software engineer': 'Software Engineer',
# }

# def assign_label(cleaned_text):
#     cleaned_text = cleaned_text.lower()
#     for keyword, label in keyword_to_label.items():
#         if keyword in cleaned_text:
#             return label
#     return 'Other'

# df['label'] = df['cleaned_text'].apply(assign_label)

# # Step 3: Prepare data
# X = df['cleaned_text']
# y = df['label']
# X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

# tokenizer = Tokenizer(num_words=5000)
# tokenizer.fit_on_texts(X_train)
# X_train_seq = tokenizer.texts_to_sequences(X_train)
# X_test_seq = tokenizer.texts_to_sequences(X_test)

# max_len = max(len(x.split()) for x in X_train)
# X_train_pad = pad_sequences(X_train_seq, maxlen=max_len)
# X_test_pad = pad_sequences(X_test_seq, maxlen=max_len)

# y_train_encoded = pd.get_dummies(y_train).values
# y_test_encoded = pd.get_dummies(y_test).values

# # Step 4: Model Training
# model = Sequential()
# model.add(Embedding(input_dim=5000, output_dim=128, input_length=max_len))
# model.add(Conv1D(filters=128, kernel_size=5, activation='relu'))
# model.add(MaxPooling1D(pool_size=2))
# model.add(Bidirectional(LSTM(100)))
# model.add(Dense(units=y_train_encoded.shape[1], activation='softmax'))

# model.compile(optimizer='adam', loss='categorical_crossentropy', metrics=['accuracy'])
# history = model.fit(X_train_pad, y_train_encoded, epochs=10, batch_size=32, validation_split=0.2)

# model.save('resume_parser_models.h5')
# === Import Libraries ===
import numpy as np
import pandas as pd
import re
import nltk
import pdfplumber
import os
from docx import Document
from tensorflow.keras.models import Sequential
from tensorflow.keras.layers import Embedding, Conv1D, MaxPooling1D, LSTM, Dense, Bidirectional
from tensorflow.keras.preprocessing.text import Tokenizer
from tensorflow.keras.preprocessing.sequence import pad_sequences
from sklearn.model_selection import train_test_split
from collections import Counter
from nltk.corpus import stopwords
from nltk.stem import PorterStemmer

nltk.download('stopwords')

# === Step 1: Extract Text ===
def extract_text_from_pdf(pdf_path):
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

def extract_text_from_docx(docx_path):
    text = ""
    doc = Document(docx_path)
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

# === Step 2: Load and Clean Resumes ===
resume_folder = r'H:\AI-Resume-Parser-main\Resumes'
resume_texts = []

for filename in os.listdir(resume_folder):
    file_path = os.path.join(resume_folder, filename)
    if filename.endswith('.pdf'):
        text = extract_text_from_pdf(file_path)
    elif filename.endswith('.docx'):
        text = extract_text_from_docx(file_path)
    else:
        continue
    resume_texts.append({'filename': filename, 'text': text})

df = pd.DataFrame(resume_texts)

# === Step 3: Text Preprocessing ===
def clean_text(text):
    text = re.sub(r'[^a-zA-Z\s]', '', text)
    text = text.lower().strip()
    return text

stop_words = set(stopwords.words('english'))
ps = PorterStemmer()

def preprocess(text):
    cleaned = clean_text(text)
    tokens = [word for word in cleaned.split() if word not in stop_words]
    stemmed = [ps.stem(word) for word in tokens]
    return ' '.join(stemmed)

df['cleaned_text'] = df['text'].apply(preprocess)

# === Step 4: Assign Labels ===
keyword_to_label = {
    'manager': 'Manager',
    'sales manager': 'Sales Manager',
    'project manager': 'Project Manager',
    'analyst': 'Analyst',
    'software engineer': 'Software Engineer',
}

def assign_label(cleaned_text):
    for keyword, label in keyword_to_label.items():
        if keyword in cleaned_text:
            return label
    return 'Other'

df['label'] = df['cleaned_text'].apply(assign_label)

# === Step 5: Prepare Data for Training ===
X = df['cleaned_text']
y = df['label']

X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

tokenizer = Tokenizer(num_words=5000)
tokenizer.fit_on_texts(X_train)

X_train_seq = tokenizer.texts_to_sequences(X_train)
X_test_seq = tokenizer.texts_to_sequences(X_test)

# Use fixed max_len = 1149 for consistency
max_len = max(len(x) for x in X_train_seq)
if max_len < 1149:
    max_len = 1149

X_train_pad = pad_sequences(X_train_seq, maxlen=max_len)
X_test_pad = pad_sequences(X_test_seq, maxlen=max_len)

# Convert labels to one-hot encoding
y_train_encoded = pd.get_dummies(y_train).values
y_test_encoded = pd.get_dummies(y_test).values

# === Step 6: Build and Train Model ===
model = Sequential()
model.add(Embedding(input_dim=5000, output_dim=128, input_length=max_len))
model.add(Conv1D(filters=128, kernel_size=5, activation='relu'))
model.add(MaxPooling1D(pool_size=2))
model.add(Bidirectional(LSTM(100)))
model.add(Dense(units=y_train_encoded.shape[1], activation='softmax'))

model.compile(optimizer='adam', loss='categorical_crossentropy', metrics=['accuracy'])

history = model.fit(X_train_pad, y_train_encoded, epochs=10, batch_size=32, validation_split=0.2)

# === Save Model ===
model.save('resume_parser_model2.h5')
